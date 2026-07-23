from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import re

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / 'APP_SPECIFICATION.md'
OUT = ROOT / 'IKUMEE_アプリ仕様書.docx'
BLUE = '0877E8'; DARK = '172033'; MUTED = '657083'; LIGHT = 'EAF5FF'; LINE = 'DCE5EF'; WHITE = 'FFFFFF'

def font(run, size=10.5, bold=False, color=DARK, name='Yu Gothic'):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), name)
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); run._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def margins(cell, top=90, start=120, bottom=90, end=120):
    tc=cell._tc.get_or_add_tcPr(); mar=tc.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); tc.append(mar)
    for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node=mar.find(qn('w:'+tag))
        if node is None: node=OxmlElement('w:'+tag); mar.append(node)
        node.set(qn('w:w'),str(val)); node.set(qn('w:type'),'dxa')

def set_cell_width(cell, dxa):
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(dxa)); tcW.set(qn('w:type'),'dxa')

def set_table_geometry(table, widths):
    table.autofit=False; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    tblPr=table._tbl.tblPr; tblW=tblPr.find(qn('w:tblW')); tblW.set(qn('w:w'),str(sum(widths))); tblW.set(qn('w:type'),'dxa')
    ind=OxmlElement('w:tblInd'); ind.set(qn('w:w'),'120'); ind.set(qn('w:type'),'dxa'); tblPr.append(ind)
    grid=table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(w)); grid.append(col)
    for row in table.rows:
        for i,c in enumerate(row.cells): set_cell_width(c,widths[i]); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def page_num(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=paragraph.add_run('Page '); font(r,8,color=MUTED)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); paragraph._p.append(fld)

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=Inches(.78); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(.82); sec.right_margin=Inches(.82); sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Yu Gothic'; normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(DARK); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Yu Gothic'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.18
for name,size,before,after,color in [('Heading 1',16,15,7,BLUE),('Heading 2',13,11,5,BLUE),('Heading 3',11.5,8,4,DARK)]:
    s=styles[name]; s.font.name='Yu Gothic'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s._element.rPr.rFonts.set(qn('w:eastAsia'),'Yu Gothic'); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True

header=sec.header.paragraphs[0]; header.text='IKUMEE  |  APP SPECIFICATION'; header.alignment=WD_ALIGN_PARAGRAPH.LEFT; font(header.runs[0],8.5,True,BLUE)
footer=sec.footer.paragraphs[0]; page_num(footer)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(42); r=p.add_run('IKUMEE'); font(r,14,True,BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(58); p.paragraph_format.space_after=Pt(5); r=p.add_run('アプリ仕様書'); font(r,29,True,DARK)
p=doc.add_paragraph(); r=p.add_run('移動支援マッチングアプリ「イクミー」'); font(r,15,False,BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(20); p.paragraph_format.space_after=Pt(28); r=p.add_run('移動に、やさしい出会いを。'); font(r,12,True,MUTED)
tbl=doc.add_table(rows=4,cols=2); set_table_geometry(tbl,[1900,7460])
meta=[('文書版数','1.0'),('作成日','2026年7月23日'),('成果物','提案・紹介動画用 静的Webプロトタイプ'),('公開方式','GitHub Pages')]
for row,(a,b) in zip(tbl.rows,meta):
    row.cells[0].text=a; row.cells[1].text=b; shade(row.cells[0],LIGHT)
    for i,c in enumerate(row.cells):
        for run in c.paragraphs[0].runs: font(run,9.5,i==0,BLUE if i==0 else DARK)
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(54); r=p.add_run('本書は授業課題における提案仕様を定義するものであり、実運用サービスのセキュリティ・法務要件を保証するものではありません。'); font(r,8.5,False,MUTED)
doc.add_page_break()

lines=SRC.read_text(encoding='utf-8').splitlines(); i=0
while i<len(lines):
    line=lines[i].strip()
    if not line or line.startswith('# IKUMEE') or line.startswith('**文書版数') or line.startswith('**作成日') or line.startswith('**対象:'):
        i+=1; continue
    if line.startswith('```'):
        block=[]; i+=1
        while i<len(lines) and not lines[i].strip().startswith('```'): block.append(lines[i]); i+=1
        p=doc.add_paragraph(); p.style=styles['Normal']; p.paragraph_format.left_indent=Inches(.18); p.paragraph_format.right_indent=Inches(.18)
        for j,t in enumerate(block):
            r=p.add_run(t+('' if j==len(block)-1 else '\n')); font(r,8.6,False,DARK,'Consolas')
        i+=1; continue
    if line.startswith('|'):
        rows=[]
        while i<len(lines) and lines[i].strip().startswith('|'): rows.append([x.strip() for x in lines[i].strip().strip('|').split('|')]); i+=1
        rows=[r for r in rows if not all(re.fullmatch(r'-+',x or '') for x in r)]
        cols=len(rows[0]); table=doc.add_table(rows=len(rows),cols=cols); table.style='Table Grid'
        if cols==2: widths=[2400,6960]
        elif cols==3: widths=[1000,2200,6160]
        else: widths=[9360//cols]*cols
        set_table_geometry(table,widths)
        for ri,row in enumerate(rows):
            for ci,text in enumerate(row):
                cell=table.cell(ri,ci); cell.text=text.replace('`','');
                if ri==0: shade(cell,BLUE)
                elif ri%2==0: shade(cell,'F7FAFD')
                for run in cell.paragraphs[0].runs: font(run,8.8,ri==0,WHITE if ri==0 else DARK)
        doc.add_paragraph().paragraph_format.space_after=Pt(1); continue
    if line.startswith('### '): p=doc.add_paragraph(line[4:],style='Heading 3')
    elif line.startswith('## '): p=doc.add_paragraph(line[3:],style='Heading 1')
    elif re.match(r'^\d+\. ',line):
        p=doc.add_paragraph(style='List Number'); p.paragraph_format.left_indent=Inches(.32); p.paragraph_format.first_line_indent=Inches(-.18); r=p.add_run(re.sub(r'^\d+\. ','',line)); font(r)
    elif line.startswith('- '):
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.left_indent=Inches(.32); p.paragraph_format.first_line_indent=Inches(-.18); r=p.add_run(line[2:]); font(r)
    else:
        p=doc.add_paragraph();
        # basic inline bold/code cleanup
        text=line.replace('  ',' ').replace('`','')
        r=p.add_run(text); font(r)
    i+=1

# keep rows together where possible, repeat headers
for table in doc.tables:
    if len(table.rows)>1:
        trPr=table.rows[0]._tr.get_or_add_trPr(); repeat=OxmlElement('w:tblHeader'); repeat.set(qn('w:val'),'true'); trPr.append(repeat)

doc.core_properties.title='IKUMEE アプリ仕様書'; doc.core_properties.subject='移動支援マッチングアプリの仕様'; doc.core_properties.author='IKUMEE Project'
doc.save(OUT)
print(OUT)
